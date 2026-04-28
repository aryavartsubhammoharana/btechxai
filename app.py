import json
import os
import re
import sys
import base64
import binascii
import ipaddress
import io
import shutil
import socket
import time
from html import unescape
from datetime import datetime
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from flask import Flask, render_template, request, jsonify
from flask_cors import CORS
from sarvamai import SarvamAI

try:
    import requests
except ImportError:
    requests = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from newspaper import Article
except ImportError:
    Article = None

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    from selenium.webdriver.common.by import By
except ImportError:
    webdriver = None
    ChromeOptions = None
    By = None

try:
    import easyocr
    import numpy as np
except ImportError:
    easyocr = None
    np = None

try:
    import torch
except ImportError:
    torch = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
except ImportError:
    Image = None
    ImageEnhance = None
    ImageFilter = None
    ImageOps = None

# ─ INITIALIZATION ─
app = Flask(__name__, template_folder='templates', static_folder='.', static_url_path='')
CORS(app)  # Enables cross-origin requests from frontend


API_KEY = os.environ.get("SARVAM_API_KEY")
client = SarvamAI(api_subscription_key=API_KEY)

# ─ CHAT STORAGE CONFIGURATION ─
CHATS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chats')
os.makedirs(CHATS_DIR, exist_ok=True)

def get_chat_storage_path(chat_id: str) -> str:
    """Get the file path for a chat storage."""
    return os.path.join(CHATS_DIR, f"{chat_id}.json")

def save_chat(chat_id: str, messages: list, metadata: dict = None) -> bool:
    """Save chat messages to a JSON file."""
    try:
        chat_data = {
            "chat_id": chat_id,
            "messages": messages,
            "metadata": metadata or {},
            "updated_at": get_timestamp()
        }
        with open(get_chat_storage_path(chat_id), 'w', encoding='utf-8') as f:
            json.dump(chat_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        log_message("ERROR", f"Failed to save chat {chat_id}: {str(e)}")
        return False

def load_chat(chat_id: str) -> dict | None:
    """Load chat messages from a JSON file."""
    try:
        path = get_chat_storage_path(chat_id)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return None
    except Exception as e:
        log_message("ERROR", f"Failed to load chat {chat_id}: {str(e)}")
        return None

def delete_chat(chat_id: str) -> bool:
    """Delete a chat from storage."""
    try:
        path = get_chat_storage_path(chat_id)
        if os.path.exists(path):
            os.remove(path)
        return True
    except Exception as e:
        log_message("ERROR", f"Failed to delete chat {chat_id}: {str(e)}")
        return False

def list_all_chats() -> list:
    """List all saved chats."""
    chats = []
    try:
        for filename in os.listdir(CHATS_DIR):
            if filename.endswith('.json'):
                chat_id = filename[:-5]
                chat_data = load_chat(chat_id)
                if chat_data:
                    chats.append({
                        "chat_id": chat_id,
                        "title": chat_data.get("metadata", {}).get("title", "Untitled Chat"),
                        "message_count": len(chat_data.get("messages", [])),
                        "updated_at": chat_data.get("updated_at", "")
                    })
    except Exception as e:
        log_message("ERROR", f"Failed to list chats: {str(e)}")
    
    # Sort by updated time, newest first
    chats.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    return chats

def generate_chat_id() -> str:
    """Generate a unique chat ID."""
    import uuid
    return str(uuid.uuid4())

# ─ SYSTEM CONFIGURATION ─
SYSTEM_PROMPT = {
    "role": "system",
    "content": (
        "You are BTechX AI — an intelligent assistant with deep knowledge of Indian culture, "
        "history, languages, science, technology, computer science, engineering, and academics. "
        "Respond like a thoughtful modern LLM: natural, helpful, and conversational. "
        "Match the user's language and tone. If the user writes in Hindi, answer in Hindi. "
        "If the user writes in English, answer in English. If the user mixes both, you may reply "
        "in natural Hinglish. Keep answers concise by default and expand only when the user asks "
        "for detail. Avoid sounding like a speech or essay unless explicitly requested. "
        "Prefer 1-3 short paragraphs. When useful, break the answer into short points or steps. "
        "For technical questions, include compact examples or code snippets when helpful."
        "Do not force every answer into one long paragraph. "
        "If website content is included in the prompt context, use it carefully and mention source links when helpful. "
        "If live website content is unavailable, be honest instead of pretending to have browsed. "
        "Answer all the Answer in 750 words, if user ask more detailed answer then give the response in between 850 words. "
        ""
    )
}

# ─ HELPER FUNCTIONS ─
SUPPORTED_FORMAT_OPTIONS = {
    "bold": "bold text with **text**",
    "italic": "italic text with *text*",
    "strikethrough": "strikethrough text with ~~text~~",
    "monospace": "inline monospace/code with `text`",
}
SUPPORTED_IMAGE_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp"}
SUPPORTED_DOCUMENT_TYPES = {
    "application/pdf",
    "text/plain",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_ATTACHMENT_TYPES = SUPPORTED_IMAGE_TYPES | SUPPORTED_DOCUMENT_TYPES
MAX_IMAGE_BYTES = 100 * 1024 * 1024
MAX_DOCUMENT_BYTES = 50 * 1024 * 1024 #50GB
OCR_READER = None
TESSERACT_AVAILABLE = pytesseract is not None and shutil.which("tesseract") is not None
GPU_AVAILABLE = torch is not None and torch.cuda.is_available()
URL_PATTERN = re.compile(r'(?P<url>(?:https?://|www\.)[^\s<>"\'`]+)', re.IGNORECASE)
MAX_URL_CONTEXTS = 3
MAX_FETCH_CHARS = 12000
ARTICLE_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0 Safari/537.36 BTechXAI/2.0"
)
ARTICLE_REQUEST_TIMEOUT = 12
SELENIUM_PAGE_LOAD_TIMEOUT = 18
MIN_ARTICLE_CHARS = 500

def get_timestamp():
    """Return current timestamp for logging."""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def log_message(level, message):
    """Log messages with timestamp."""
    prefix = f"[{get_timestamp()}] [{level}]"
    print(f"{prefix} {message}")
    sys.stdout.flush()

def clean_response(text: str) -> str:
    """Normalize model output while preserving readable structure."""
    if not text:
        return ""

    text = text.replace('\r\n', '\n').replace('\r', '\n').strip()
    text = re.sub(r'([?!])\s+(For example[:,]?)', r'\1\n\n\2', text, flags=re.IGNORECASE)
    text = re.sub(r'([.!?])\s+(Example[:,]?)', r'\1\n\n\2', text, flags=re.IGNORECASE)
    text = re.sub(r'([.!?])\s+(Here are|Try this|You can also|For instance[:,]?)', r'\1\n\n\2', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+-\s+', '\n- ', text)
    text = re.sub(r'\n(- .+?)\s+-\s+', r'\n\1\n- ', text)

    # Trim overly large empty gaps but keep paragraph breaks intact.
    while '\n\n\n' in text:
        text = text.replace('\n\n\n', '\n\n')

    return text

def extract_urls(text: str) -> list[str]:
    """Return normalized HTTP(S) URLs from message text."""
    if not text:
        return []

    urls = []
    seen = set()
    for match in URL_PATTERN.finditer(text):
        candidate = match.group("url").rstrip(".,;:!?)]}")
        normalized = candidate if candidate.lower().startswith(("http://", "https://")) else f"https://{candidate}"
        if normalized not in seen:
            seen.add(normalized)
            urls.append(normalized)
    return urls

def html_to_text(html: str) -> str:
    """Strip HTML to compact readable text."""
    if not html:
        return ""

    html = re.sub(r"(?is)<script\b.*?</script>", " ", html)
    html = re.sub(r"(?is)<style\b.*?</style>", " ", html)

    title_match = re.search(r"(?is)<title[^>]*>(.*?)</title>", html)
    title_text = ""
    if title_match:
        title_text = unescape(re.sub(r"\s+", " ", title_match.group(1))).strip()

    body_text = re.sub(r"(?s)<[^>]+>", " ", html)
    body_text = unescape(body_text)
    body_text = re.sub(r"\s+", " ", body_text).strip()

    if title_text and not body_text.lower().startswith(title_text.lower()):
        body_text = f"{title_text}. {body_text}"

    return body_text[:MAX_FETCH_CHARS]

def compact_text(text: str, max_chars: int = MAX_FETCH_CHARS) -> str:
    """Normalize extracted article text into compact readable context."""
    if not text:
        return ""

    text = unescape(str(text))
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    return text[:max_chars]

def is_public_url(url: str) -> tuple[bool, str]:
    """Allow only public HTTP(S) URLs to avoid fetching local/private services."""
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        return False, "Unsupported URL scheme."

    hostname = parsed.hostname
    if not hostname:
        return False, "URL is missing a hostname."

    if hostname.lower() in {"localhost", "localhost.localdomain"}:
        return False, "Localhost URLs are not allowed."

    try:
        addresses = socket.getaddrinfo(hostname, None)
    except socket.gaierror:
        return False, "Could not resolve website hostname."

    for address in addresses:
        ip_value = address[4][0]
        try:
            ip = ipaddress.ip_address(ip_value)
        except ValueError:
            continue

        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
            return False, "Private or local network URLs are not allowed."

    return True, ""

def build_article_result(url: str, method: str, content: str, title: str = "", metadata: dict | None = None) -> dict:
    """Create a consistent fetch result for website context blocks."""
    return {
        "url": url,
        "status": "ok",
        "method": method,
        "title": compact_text(title, 300),
        "content": compact_text(content),
        "metadata": metadata or {},
    }

def extract_with_newspaper(url: str, html: str | None = None) -> dict | None:
    """Extract article text using newspaper3k when available."""
    if Article is None:
        return None

    try:
        article = Article(url, browser_user_agent=ARTICLE_USER_AGENT)
        if html:
            article.set_html(html)
        else:
            article.download()
        article.parse()

        content = compact_text(article.text)
        if len(content) < MIN_ARTICLE_CHARS:
            return None

        metadata = {
            "authors": ", ".join(article.authors[:5]) if article.authors else "",
            "publish_date": str(article.publish_date) if article.publish_date else "",
        }
        return build_article_result(url, "newspaper3k", content, article.title, metadata)
    except Exception as newspaper_error:
        log_message("WARNING", f"newspaper3k extraction failed for {url}: {str(newspaper_error)}")
        return None

def extract_with_beautifulsoup(url: str, html: str) -> dict | None:
    """Extract readable article text using BeautifulSoup."""
    if BeautifulSoup is None or not html:
        return None

    try:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "iframe", "form", "nav", "footer", "header"]):
            tag.decompose()

        title = ""
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

        candidates = []
        for selector in ("article", "main", "[role='main']", ".article", ".post", ".entry-content", ".content"):
            for node in soup.select(selector):
                text = compact_text(node.get_text("\n"))
                if text:
                    candidates.append(text)

        if not candidates:
            paragraphs = [compact_text(p.get_text(" ")) for p in soup.find_all(["p", "li"])]
            candidates = [paragraph for paragraph in paragraphs if len(paragraph) > 40]

        content = compact_text("\n\n".join(candidates))
        if len(content) < 200:
            content = compact_text(soup.get_text("\n"))
        if not content:
            return None

        return build_article_result(url, "requests + BeautifulSoup", content, title)
    except Exception as soup_error:
        log_message("WARNING", f"BeautifulSoup extraction failed for {url}: {str(soup_error)}")
        return None

def fetch_html_with_requests(url: str) -> tuple[str, str]:
    """Fetch raw website HTML using requests."""
    if requests is None:
        return "", "requests is not installed."

    response = requests.get(
        url,
        headers={
            "User-Agent": ARTICLE_USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,text/plain;q=0.8,*/*;q=0.7",
            "Accept-Language": "en-US,en;q=0.9,hi;q=0.8",
        },
        timeout=ARTICLE_REQUEST_TIMEOUT,
        allow_redirects=True,
    )
    response.raise_for_status()

    content_type = response.headers.get("Content-Type", "").lower()
    if "text/html" not in content_type and "text/plain" not in content_type and content_type:
        return "", f"Skipped non-text response ({content_type})."

    response.encoding = response.encoding or response.apparent_encoding or "utf-8"
    return response.text, ""

def extract_with_selenium(url: str) -> dict | None:
    """Render JavaScript-heavy pages with Selenium as a final fallback."""
    if webdriver is None or ChromeOptions is None:
        return None

    driver = None
    try:
        options = ChromeOptions()
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--window-size=1365,900")
        options.add_argument(f"--user-agent={ARTICLE_USER_AGENT}")

        driver = webdriver.Chrome(options=options)
        driver.set_page_load_timeout(SELENIUM_PAGE_LOAD_TIMEOUT)
        driver.get(url)
        time.sleep(2)

        title = driver.title or ""
        html = driver.page_source or ""
        result = extract_with_newspaper(url, html) or extract_with_beautifulsoup(url, html)
        if result:
            result["method"] = f"selenium + {result['method']}"
            if title and not result.get("title"):
                result["title"] = compact_text(title, 300)
            return result

        body = driver.find_element(By.TAG_NAME, "body").text if By is not None else ""
        if body:
            return build_article_result(url, "selenium", body, title)
    except Exception as selenium_error:
        log_message("WARNING", f"Selenium extraction failed for {url}: {str(selenium_error)}")
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass

    return None

def fetch_website_context(url: str) -> dict:
    """Fetch readable article/page text from a user-provided public website."""
    is_allowed, block_reason = is_public_url(url)
    if not is_allowed:
        return {
            "url": url,
            "status": "blocked",
            "method": "security-check",
            "title": "",
            "content": block_reason,
            "metadata": {},
        }

    try:
        newspaper_result = extract_with_newspaper(url)
        if newspaper_result:
            return newspaper_result

        html, fetch_note = fetch_html_with_requests(url)
        if html:
            soup_result = extract_with_beautifulsoup(url, html)
            if soup_result:
                return soup_result

            newspaper_from_html = extract_with_newspaper(url, html)
            if newspaper_from_html:
                return newspaper_from_html

            fallback_text = html_to_text(html)
            if fallback_text:
                return build_article_result(url, "requests fallback", fallback_text)
        elif fetch_note:
            log_message("WARNING", f"requests fetch note for {url}: {fetch_note}")

        selenium_result = extract_with_selenium(url)
        if selenium_result:
            return selenium_result

        if fetch_note:
            return {
                "url": url,
                "status": "unsupported",
                "method": "requests",
                "title": "",
                "content": fetch_note,
                "metadata": {},
            }

        return {
            "url": url,
            "status": "empty",
            "method": "article-readers",
            "title": "",
            "content": "No readable text found on this page.",
            "metadata": {},
        }
    except HTTPError as http_error:
        return {"url": url, "status": "error", "method": "urllib", "title": "", "content": f"Website returned HTTP {http_error.code}.", "metadata": {}}
    except URLError as url_error:
        return {"url": url, "status": "error", "method": "urllib", "title": "", "content": f"Could not reach website: {url_error.reason}.", "metadata": {}}
    except Exception as fetch_error:
        selenium_result = extract_with_selenium(url)
        if selenium_result:
            return selenium_result

        log_message("WARNING", f"Website fetch failed for {url}: {str(fetch_error)}")
        return {"url": url, "status": "error", "method": "article-readers", "title": "", "content": f"Could not fetch website content: {str(fetch_error)}.", "metadata": {}}

def build_website_context(message: str) -> str:
    """Build supplemental website text from URLs included in the prompt."""
    urls = extract_urls(message)[:MAX_URL_CONTEXTS]
    if not urls:
        return ""

    context_blocks = []
    for url in urls:
        result = fetch_website_context(url)
        lines = [
            "[Website context]",
            f"Source: {result['url']}",
            f"Status: {result.get('status', 'unknown')}",
            f"Reader: {result.get('method', 'unknown')}",
        ]
        if result.get("title"):
            lines.append(f"Title: {result['title']}")

        metadata = result.get("metadata") or {}
        if metadata.get("authors"):
            lines.append(f"Authors: {metadata['authors']}")
        if metadata.get("publish_date"):
            lines.append(f"Published: {metadata['publish_date']}")

        lines.append(f"Content: {result.get('content', '')}")
        context_blocks.append("\n".join(lines))

    return "\n\n".join(context_blocks)

def get_format_options(data: dict) -> dict:
    """Return user-selected response formatting preferences."""
    requested_options = data.get('formatOptions', {})
    if not isinstance(requested_options, dict):
        requested_options = {}

    return {
        option: bool(requested_options.get(option, True))
        for option in SUPPORTED_FORMAT_OPTIONS
    }

def build_format_prompt(format_options: dict) -> dict:
    """Build a system instruction that limits markdown styles in AI replies."""
    enabled = [
        description
        for option, description in SUPPORTED_FORMAT_OPTIONS.items()
        if format_options.get(option)
    ]
    disabled = [
        option.replace("_", " ")
        for option in SUPPORTED_FORMAT_OPTIONS
        if not format_options.get(option)
    ]

    if enabled:
        content = (
            "Response formatting preference: Use only these Markdown styles when useful: "
            f"{', '.join(enabled)}. "
        )
    else:
        content = (
            "Response formatting preference: Reply in plain text only. "
            "Do not use Markdown formatting symbols for bold, italic, strikethrough, or monospace/code. "
        )

    if disabled:
        content += f"Do not use these disabled styles: {', '.join(disabled)}. "

    content += "Keep the answer natural; formatting is optional, not required."
    return {"role": "system", "content": content}

def apply_format_preferences(text: str, format_options: dict) -> str:
    """Remove disabled markdown styles if the model still returns them."""
    if not text:
        return ""

    if not format_options.get("monospace", True):
        text = re.sub(r'`([^`\n]+)`', r'\1', text)
    if not format_options.get("strikethrough", True):
        text = re.sub(r'~~(.+?)~~', r'\1', text, flags=re.DOTALL)
    if not format_options.get("bold", True):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'__(.+?)__', r'\1', text, flags=re.DOTALL)
    if not format_options.get("italic", True):
        text = re.sub(r'(?<!\*)\*(?!\*)([^*\n]+?)(?<!\*)\*(?!\*)', r'\1', text)
        text = re.sub(r'(?<!_)_(?!_)([^_\n]+?)(?<!_)_(?!_)', r'\1', text)

    return text

def get_ocr_reader():
    """Create the OCR reader only when an image is uploaded."""
    global OCR_READER

    if OCR_READER is not None:
        return OCR_READER
    if easyocr is None:
        return None

    try:
        OCR_READER = easyocr.Reader(["en", "hi"], gpu=GPU_AVAILABLE)
        log_message("INFO", f"EasyOCR initialized with {'GPU' if GPU_AVAILABLE else 'CPU'}")
        return OCR_READER
    except Exception as ocr_error:
        log_message("WARNING", f"OCR initialization failed: {str(ocr_error)}")
        return None

def parse_uploaded_attachment(attachment_payload: dict) -> tuple[dict | None, str]:
    """Validate and decode a base64 attachment sent from the browser."""
    if not attachment_payload:
        return None, ""
    if not isinstance(attachment_payload, dict):
        return None, "Attachment payload must be an object"

    attachment_type = attachment_payload.get("type", "")
    attachment_name = attachment_payload.get("name", "uploaded-file")
    attachment_data = attachment_payload.get("data", "")
    is_image = attachment_type in SUPPORTED_IMAGE_TYPES

    if attachment_type not in SUPPORTED_ATTACHMENT_TYPES:
        return None, "Only PNG, JPG, JPEG, WEBP, PDF, TXT, DOC, and DOCX files are supported"
    if not attachment_data:
        return None, "Attachment data is missing"

    if "," in attachment_data:
        attachment_data = attachment_data.split(",", 1)[1]

    try:
        attachment_bytes = base64.b64decode(attachment_data, validate=True)
    except (binascii.Error, ValueError):
        return None, "Attachment data is not valid base64"

    max_bytes = MAX_IMAGE_BYTES if is_image else MAX_DOCUMENT_BYTES
    if len(attachment_bytes) > max_bytes:
        max_mb = round(max_bytes / (1024 * 1024))
        return None, f"Attachment must be {max_mb} MB or smaller"

    if is_image and len(attachment_bytes) > MAX_IMAGE_BYTES:
        return None, f"Image must be 100 MB or smaller"

    return {
        "name": str(attachment_name)[:120],
        "type": attachment_type,
        "size": len(attachment_bytes),
        "bytes": attachment_bytes,
        "kind": "image" if is_image else "document",
    }, ""

def parse_uploaded_image(image_payload: dict) -> tuple[dict | None, str]:
    """Backward-compatible wrapper for older image payloads."""
    return parse_uploaded_attachment(image_payload)

def build_ocr_variants(image):
    """Create OCR-friendly versions of an uploaded image."""
    if image.mode != "RGB":
        image = image.convert("RGB")

    variants = [image]
    width, height = image.size
    upscale_factor = 2 if max(width, height) < 2200 else 1

    if upscale_factor > 1:
        variants.append(image.resize((width * upscale_factor, height * upscale_factor), Image.Resampling.LANCZOS))

    if ImageOps and ImageEnhance and ImageFilter:
        grayscale = ImageOps.grayscale(variants[-1])
        contrast = ImageEnhance.Contrast(grayscale).enhance(2.2)
        sharpened = contrast.filter(ImageFilter.SHARPEN)
        inverted = ImageOps.invert(contrast)
        binary = contrast.point(lambda pixel: 255 if pixel > 165 else 0)
        variants.extend([grayscale, contrast, sharpened, inverted, binary])

    return variants

def clean_ocr_text(text: str) -> str:
    """Normalize noisy OCR output."""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def ocr_text_score(text: str) -> int:
    """Score OCR output by useful readable characters."""
    if not text:
        return 0
    return len(re.findall(r"[A-Za-z0-9]", text))

def extract_text_from_uploaded_image(image_info: dict) -> str:
    """Extract readable text from an uploaded image with OCR."""
    if not image_info:
        return ""
    if Image is None:
        return ""

    try:
        image = Image.open(io.BytesIO(image_info["bytes"]))
        best_text = ""
        best_score = 0
        reader = get_ocr_reader()

        for candidate in build_ocr_variants(image):
            if reader is not None:
                for paragraph_mode in (False, True):
                    try:
                        ocr_input = np.array(candidate) if np is not None else candidate
                        results = reader.readtext(ocr_input, detail=0, paragraph=paragraph_mode)
                        candidate_text = clean_ocr_text("\n".join(results))
                        candidate_score = ocr_text_score(candidate_text)
                        if candidate_score > best_score:
                            best_text = candidate_text
                            best_score = candidate_score
                    except Exception as variant_error:
                        log_message("WARNING", f"EasyOCR variant failed: {str(variant_error)}")

            if TESSERACT_AVAILABLE:
                for config in ("--psm 6", "--psm 11"):
                    try:
                        candidate_text = clean_ocr_text(
                            pytesseract.image_to_string(candidate, config=config)
                        )
                        candidate_score = ocr_text_score(candidate_text)
                        if candidate_score > best_score:
                            best_text = candidate_text
                            best_score = candidate_score
                    except Exception as tesseract_error:
                        log_message("WARNING", f"Tesseract OCR failed: {str(tesseract_error)}")

        return best_text
    except Exception as ocr_error:
        log_message("WARNING", f"Image OCR failed: {str(ocr_error)}")
        return ""

def extract_text_from_pdf(attachment_info: dict) -> str:
    """Extract text from PDF pages."""
    if fitz is None:
        return "PDF extraction unavailable: PyMuPDF is not installed."

    try:
        text_parts = []
        with fitz.open(stream=attachment_info["bytes"], filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text("text"))
        return "\n".join(text_parts).strip()
    except Exception as pdf_error:
        log_message("WARNING", f"PDF extraction failed: {str(pdf_error)}")
        return "Could not extract readable text from this PDF."

def extract_text_from_txt(attachment_info: dict) -> str:
    """Decode a plain text attachment."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return attachment_info["bytes"].decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return "Could not decode this text file."

def extract_text_from_docx(attachment_info: dict) -> str:
    """Extract text from a DOCX attachment."""
    if Document is None:
        return "DOCX extraction unavailable: python-docx is not installed."

    try:
        document = Document(io.BytesIO(attachment_info["bytes"]))
        text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        return "\n".join(text_parts).strip()
    except Exception as docx_error:
        log_message("WARNING", f"DOCX extraction failed: {str(docx_error)}")
        return "Could not extract readable text from this DOCX file."

def extract_text_from_doc(attachment_info: dict) -> str:
    """Best-effort extraction for legacy DOC files."""
    raw_text = attachment_info["bytes"].decode("latin-1", errors="ignore")
    readable_chunks = re.findall(r"[A-Za-z0-9\s.,;:!?()\[\]{}'\"/@#%&+\-=]{4,}", raw_text)
    cleaned_text = "\n".join(chunk.strip() for chunk in readable_chunks if chunk.strip())
    cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text).strip()
    return cleaned_text or "Could not extract readable text from this legacy DOC file."

def extract_text_from_attachment(attachment_info: dict) -> str:
    """Extract text/OCR context from a supported attachment."""
    if not attachment_info:
        return ""

    attachment_type = attachment_info["type"]
    if attachment_type in SUPPORTED_IMAGE_TYPES:
        return extract_text_from_uploaded_image(attachment_info)
    if attachment_type == "application/pdf":
        return extract_text_from_pdf(attachment_info)
    if attachment_type == "text/plain":
        return extract_text_from_txt(attachment_info)
    if attachment_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(attachment_info)
    if attachment_type == "application/msword":
        return extract_text_from_doc(attachment_info)

    return ""

def build_user_content(message: str, attachment_info: dict | None, attachment_text: str) -> str:
    """Combine the user's text with attachment metadata/extracted text for the text-only model."""
    website_context = build_website_context(message)

    if not attachment_info:
        return "\n\n".join(part for part in [message, website_context] if part)

    attachment_label = "image" if attachment_info["kind"] == "image" else "document"
    parts = [
        message or f"Please analyze the uploaded {attachment_label}.",
        "",
        f"[Uploaded {attachment_label} context]",
        f"Name: {attachment_info['name']}",
        f"Type: {attachment_info['type']}",
        f"Size: {round(attachment_info['size'] / 1024, 1)} KB",
    ]

    if attachment_text:
        text_label = "OCR text found in image:" if attachment_info["kind"] == "image" else "Extracted document text:"
        parts.extend(["", text_label, attachment_text[:15000]])
    else:
        no_text_note = (
            "OCR text found in image: No readable text detected. The app can OCR text from images, "
            "but it cannot fully inspect visual-only details without a vision-capable model."
            if attachment_info["kind"] == "image"
            else "Extracted document text: No readable text detected."
        )
        parts.extend(["", no_text_note])

    if attachment_info["kind"] == "image":
        parts.append(
            "If the user asks about visual details not present in OCR text, be honest that only OCR/text context is available."
        )
    if website_context:
        parts.extend(["", website_context])
    return "\n".join(parts)

def validate_request(data: dict) -> tuple[bool, str]:
    """Validate incoming request data."""
    if not data:
        return False, "Request body is empty"
    
    message = data.get('message', '').strip()
    attachment_payload = data.get('attachment') or data.get('image')
    if not message and not attachment_payload:
        return False, "Message or attachment is required"
    
    if len(message) > 5000:
        return False, "Message exceeds maximum length (5000 characters)"
    
    history = data.get('history', [])
    if not isinstance(history, list):
        return False, "History must be a list"

    format_options = data.get('formatOptions', {})
    if format_options and not isinstance(format_options, dict):
        return False, "Format options must be an object"

    if attachment_payload:
        _, attachment_error = parse_uploaded_attachment(attachment_payload)
        if attachment_error:
            return False, attachment_error
    
    return True, ""

# ─ API ENDPOINTS ─
@app.route('/')
def index():
    return render_template('index.html')
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'service': 'BTechX AI Server',
        'timestamp': get_timestamp()
    }), 200

@app.route('/chat', methods=['POST'])
def chat():
    """Main chat endpoint."""
    try:
        data = request.get_json()
        
        # Validate request
        is_valid, error_msg = validate_request(data)
        if not is_valid:
            log_message("WARNING", f"Invalid request: {error_msg}")
            return jsonify({'error': error_msg}), 400
        
        user_message = data.get('message', '').strip()
        client_history = data.get('history', [])
        format_options = get_format_options(data)
        attachment_info, attachment_error = parse_uploaded_attachment(data.get('attachment') or data.get('image'))
        if attachment_error:
            return jsonify({'error': attachment_error}), 400
        attachment_text = extract_text_from_attachment(attachment_info) if attachment_info else ""
        user_content = build_user_content(user_message, attachment_info, attachment_text)
        
        log_message("INFO", f"User message: {user_message[:50]}...")
        
        # Build conversation messages array
        messages = [SYSTEM_PROMPT, build_format_prompt(format_options)]
        
        # Add conversation history (excluding the current user message already in frontend)
        for turn in client_history[:-1]:
            if turn.get('role') in ('user', 'assistant'):
                messages.append({
                    'role': turn['role'],
                    'content': turn['content']
                })
        
        # Append current user message
        messages.append({'role': 'user', 'content': user_content})
        
        try:
            log_message("INFO", "Calling Sarvam AI API...")
            
            response = client.chat.completions(
                messages=messages,
                model="sarvam-30b",
                temperature=0.8,
                reasoning_effort="medium"
            )
            
            assistant_text = response.choices[0].message.content
            cleaned_response = clean_response(assistant_text)
            cleaned_response = apply_format_preferences(cleaned_response, format_options)
            
            log_message("INFO", f"Response generated: {cleaned_response[:50]}...")
            
            return jsonify({
                'response': cleaned_response,
                'success': True,
                'timestamp': get_timestamp()
            }), 200
        
        except AttributeError as ae:
            log_message("ERROR", f"API response format error: {str(ae)}")
            return jsonify({
                'error': 'Unexpected response format from AI service',
                'details': str(ae)
            }), 500
        
        except Exception as sarvam_error:
            log_message("ERROR", f"Sarvam AI error: {str(sarvam_error)}")
            return jsonify({
                'error': 'Failed to get response from AI service',
                'details': str(sarvam_error)
            }), 502
    
    except json.JSONDecodeError:
        log_message("ERROR", "Invalid JSON in request body")
        return jsonify({'error': 'Invalid JSON format'}), 400
    
    except Exception as e:
        log_message("ERROR", f"Unexpected server error: {str(e)}")
        return jsonify({
            'error': 'Internal server error',
            'details': str(e)
        }), 500

@app.route('/info', methods=['GET'])
def info():
    """Return server information."""
    return jsonify({
        'name': 'BTechX AI Server',
        'version': '2.0',
        'model': 'sarvam-30b',
        'ocr_device': 'gpu' if GPU_AVAILABLE else 'cpu',
        'endpoints': [
            '/health',
            '/info',
            '/chat',
            '/chats',
            '/chat/<chat_id>',
            '/chat/<chat_id>/messages'
        ],
        'features': [
            'Multi-turn conversation support',
            'Markdown cleaning',
            'User-controlled response formatting',
            'Article and website reading via newspaper3k, requests, BeautifulSoup, and Selenium fallback',
            'Image upload/paste with OCR context',
            'PDF, TXT, DOC, and DOCX upload support',
            'Request validation',
            'Error handling',
            'Comprehensive logging',
            'Chat history storage & memory'
        ]
    }), 200

# ─ CHAT STORAGE ENDPOINTS ─
@app.route('/chats', methods=['GET'])
def get_chats():
    """Get list of all saved chats."""
    try:
        chats = list_all_chats()
        return jsonify({
            'chats': chats,
            'count': len(chats),
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to get chats: {str(e)}")
        return jsonify({'error': 'Failed to retrieve chats'}), 500

@app.route('/chats', methods=['POST'])
def create_chat():
    """Create a new chat."""
    try:
        data = request.get_json() or {}
        chat_id = generate_chat_id()
        title = data.get('title', 'New Chat')
        
        metadata = {
            "title": title,
            "created_at": get_timestamp()
        }
        
        save_chat(chat_id, [], metadata)
        
        return jsonify({
            'chat_id': chat_id,
            'title': title,
            'success': True
        }), 201
    except Exception as e:
        log_message("ERROR", f"Failed to create chat: {str(e)}")
        return jsonify({'error': 'Failed to create chat'}), 500

@app.route('/chat/<chat_id>', methods=['GET'])
def get_chat(chat_id):
    """Get a specific chat by ID."""
    try:
        chat_data = load_chat(chat_id)
        if chat_data is None:
            return jsonify({'error': 'Chat not found'}), 404
        
        return jsonify({
            'chat_id': chat_id,
            'messages': chat_data.get('messages', []),
            'metadata': chat_data.get('metadata', {}),
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to get chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to retrieve chat'}), 500

@app.route('/chat/<chat_id>', methods=['PUT'])
def update_chat(chat_id):
    """Update chat metadata (title, etc)."""
    try:
        data = request.get_json() or {}
        chat_data = load_chat(chat_id)
        
        if chat_data is None:
            return jsonify({'error': 'Chat not found'}), 404
        
        # Update metadata
        if 'title' in data:
            chat_data['metadata']['title'] = data['title']
        
        save_chat(chat_id, chat_data.get('messages', []), chat_data.get('metadata', {}))
        
        return jsonify({
            'chat_id': chat_id,
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to update chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to update chat'}), 500

@app.route('/chat/<chat_id>', methods=['DELETE'])
def delete_chat_endpoint(chat_id):
    """Delete a specific chat."""
    try:
        if delete_chat(chat_id):
            return jsonify({
                'chat_id': chat_id,
                'success': True
            }), 200
        return jsonify({'error': 'Chat not found'}), 404
    except Exception as e:
        log_message("ERROR", f"Failed to delete chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to delete chat'}), 500

@app.route('/chat/<chat_id>/messages', methods=['POST'])
def add_message_to_chat(chat_id):
    """Add a message to an existing chat."""
    try:
        data = request.get_json() or {}
        role = data.get('role', 'user')
        content = data.get('content', '')
        
        if not content:
            return jsonify({'error': 'Message content is required'}), 400
        
        chat_data = load_chat(chat_id)
        
        if chat_data is None:
            # Create new chat if doesn't exist
            chat_data = {
                "chat_id": chat_id,
                "messages": [],
                "metadata": {"title": "Chat"},
                "updated_at": get_timestamp()
            }
        
        messages = chat_data.get('messages', [])
        messages.append({
            "role": role,
            "content": content,
            "timestamp": get_timestamp()
        })
        
        # Auto-generate title from first user message if untitled
        if chat_data.get('metadata', {}).get('title') == 'New Chat' and role == 'user':
            title = content[:50] + ('...' if len(content) > 50 else '')
            chat_data['metadata']['title'] = title
        
        save_chat(chat_id, messages, chat_data.get('metadata', {}))
        
        return jsonify({
            'chat_id': chat_id,
            'message_count': len(messages),
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to add message to chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to add message'}), 500

@app.route('/chat/<chat_id>/messages', methods=['GET'])
def get_chat_messages(chat_id):
    """Get all messages from a chat."""
    try:
        chat_data = load_chat(chat_id)
        if chat_data is None:
            return jsonify({'error': 'Chat not found'}), 404
        
        return jsonify({
            'chat_id': chat_id,
            'messages': chat_data.get('messages', []),
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to get messages for chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to retrieve messages'}), 500

@app.route('/chat/<chat_id>/messages', methods=['DELETE'])
def clear_chat_messages(chat_id):
    """Clear all messages from a chat."""
    try:
        chat_data = load_chat(chat_id)
        
        if chat_data is None:
            return jsonify({'error': 'Chat not found'}), 404
        
        save_chat(chat_id, [], chat_data.get('metadata', {}))
        
        return jsonify({
            'chat_id': chat_id,
            'success': True
        }), 200
    except Exception as e:
        log_message("ERROR", f"Failed to clear messages for chat {chat_id}: {str(e)}")
        return jsonify({'error': 'Failed to clear messages'}), 500

# ─ ERROR HANDLERS ─
@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors."""
    log_message("WARNING", f"404 error: {request.path}")
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(405)
def method_not_allowed(error):
    """Handle 405 errors."""
    log_message("WARNING", f"405 error: {request.method} {request.path}")
    return jsonify({'error': 'Method not allowed'}), 405

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors."""
    log_message("ERROR", f"500 error: {str(error)}")
    return jsonify({'error': 'Internal server error'}), 500

# ─ BEFORE/AFTER HANDLERS ─
@app.before_request
def log_request():
    """Log incoming requests."""
    log_message("INFO", f"Request: {request.method} {request.path}")

@app.after_request
def log_response(response):
    """Log response status."""
    log_message("INFO", f"Response: {response.status_code}")
    return response

port = int(os.environ.get("PORT", 5000))
# ─ APPLICATION ENTRY POINT ─
if __name__ == '__main__':
    log_message("INFO", "=" * 60)
    log_message("INFO", "BTechX AI Server v2.0")
    log_message("INFO", "Powered by Sarvam AI (sarvam-30b)")
    log_message("INFO", "=" * 60)
    log_message("INFO", f"Starting server on http://0.0.0.0:{port}")
    log_message("INFO", f"Health check: GET http://0.0.0.0:{port}/health")
    log_message("INFO", f"Chat endpoint: POST http://0.0.0.0:{port}/chat")
    log_message("INFO", "=" * 60)
    
    try:
        app.run(debug=False, host='0.0.0.0', port=port)
    except KeyboardInterrupt:
        log_message("INFO", "Server shutdown initiated")
    except Exception as e:
        log_message("ERROR", f"Server failed to start: {str(e)}")
